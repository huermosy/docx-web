from fastapi import APIRouter, HTTPException, UploadFile, File
from fastapi.responses import FileResponse
from .analyze import TASK_STORE
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib import colors
from datetime import datetime
from pathlib import Path
import uuid

router = APIRouter(prefix="/api", tags=["report"])

REPORTS_DIR = Path(__file__).parent.parent.parent / "reports"
REPORTS_DIR.mkdir(exist_ok=True)

TEMPLATES_DIR = Path(__file__).parent.parent.parent / "templates"
TEMPLATES_DIR.mkdir(exist_ok=True)


@router.post("/upload/template")
async def upload_template(file: UploadFile = File(...)):
    """上传文档模板"""
    content = await file.read()

    # 验证文件类型
    allowed = [".dotx", ".dot", ".docx", ".doc"]
    if not any(file.filename.endswith(ext) for ext in allowed):
        raise HTTPException(
            status_code=400,
            detail="Only template files (.dotx, .dot, .docx, .doc) are supported"
        )

    # 保存模板
    template_path = TEMPLATES_DIR / f"custom_{uuid.uuid4()}{Path(file.filename).suffix}"
    with open(template_path, "wb") as f:
        f.write(content)

    return {"success": True, "template_id": template_path.stem, "filename": file.filename}


@router.get("/report/{task_id}/pdf")
async def download_pdf(task_id: str):
    """下载 PDF 格式报告"""
    task = TASK_STORE.get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")

    if task.status.value != "done":
        raise HTTPException(status_code=400, detail="Analysis not complete")

    # 生成 PDF
    pdf_path = REPORTS_DIR / f"{task_id}.pdf"
    generate_pdf(pdf_path, task)

    return FileResponse(pdf_path, media_type="application/pdf", filename=f"report_{task_id}.pdf")


@router.get("/report/{task_id}/docx")
async def download_docx(task_id: str):
    """下载 Word 格式报告"""
    task = TASK_STORE.get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")

    if task.status.value != "done":
        raise HTTPException(status_code=400, detail="Analysis not complete")

    # 生成 Word
    docx_path = REPORTS_DIR / f"{task_id}.docx"
    generate_docx(docx_path, task)

    return FileResponse(
        docx_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"report_{task_id}.docx"
    )


def generate_pdf(path, task):
    """生成 PDF 报告"""
    doc = SimpleDocTemplate(str(path), pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    # 标题
    story.append(Paragraph("Word 文档质量检查报告", styles['Title']))
    story.append(Spacer(1, 20))

    # 统计
    issues = task.merged_result or []
    critical = sum(1 for i in issues if i.severity.value == 'critical')
    major = sum(1 for i in issues if i.severity.value == 'major')
    minor = sum(1 for i in issues if i.severity.value == 'minor')

    story.append(Paragraph(f"<b>检查时间:</b> {task.updated_at.strftime('%Y-%m-%d %H:%M')}", styles['Normal']))
    story.append(Paragraph(f"<b>总问题数:</b> {len(issues)}", styles['Normal']))
    story.append(Paragraph(f"<b>严重问题:</b> {critical}", styles['Normal']))
    story.append(Paragraph(f"<b>中等问题:</b> {major}", styles['Normal']))
    story.append(Paragraph(f"<b>轻微问题:</b> {minor}", styles['Normal']))
    story.append(Spacer(1, 20))

    # 问题列表
    story.append(Paragraph("<b>详细问题列表</b>", styles['Heading2']))

    for idx, issue in enumerate(issues, 1):
        story.append(Paragraph(
            f"<b>{idx}. [{issue.severity.value.upper()}] {issue.description}</b>",
            styles['Normal']
        ))
        story.append(Paragraph(f"位置: {issue.position.section} - 段落 {issue.position.paragraph}", styles['Normal']))
        story.append(Paragraph(f"建议: {issue.fixSuggestion}", styles['Normal']))
        story.append(Spacer(1, 10))

    doc.build(story)


def generate_docx(path, task):
    """生成 Word 报告"""
    from docx import Document

    doc = Document()

    # 标题
    doc.add_heading('Word 文档质量检查报告', 0)

    # 统计
    issues = task.merged_result or []
    critical = sum(1 for i in issues if i.severity.value == 'critical')
    major = sum(1 for i in issues if i.severity.value == 'major')
    minor = sum(1 for i in issues if i.severity.value == 'minor')

    doc.add_paragraph(f"检查时间: {task.updated_at.strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"总问题数: {len(issues)}")
    doc.add_paragraph(f"严重问题: {critical}")
    doc.add_paragraph(f"中等问题: {major}")
    doc.add_paragraph(f"轻微问题: {minor}")

    doc.add_heading('详细问题列表', 1)

    for idx, issue in enumerate(issues, 1):
        p = doc.add_paragraph()
        severity_text = f"[{issue.severity.value.upper()}]"
        p.add_run(f"{idx}. {severity_text} {issue.description}").bold = True
        doc.add_paragraph(f"位置: {issue.position.section} - 段落 {issue.position.paragraph}")
        doc.add_paragraph(f"建议: {issue.fixSuggestion}")
        doc.add_paragraph()

    doc.save(str(path))

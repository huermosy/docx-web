from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from typing import List, Dict, Tuple, Any
import re


class DocumentParser:
    """Word 文档解析器，用于提取文档结构和内容"""

    def __init__(self, file_path: str):
        """
        初始化解析器，加载 docx 文件

        Args:
            file_path: docx 文件路径
        """
        self.file_path = file_path
        self.document = Document(file_path)
        self._paragraphs_cache = None
        self._headings_cache = None

    def extract_paragraphs(self) -> List[Dict[str, Any]]:
        """
        返回所有段落列表

        Returns:
            List[Dict]: 每个段落包含 text, style_name, is_heading
        """
        paragraphs = []
        for para in self.document.paragraphs:
            text = para.text.strip()
            style_name = para.style.name if para.style else "Normal"
            is_heading = self._is_heading_style(style_name)

            paragraphs.append({
                "text": text,
                "style_name": style_name,
                "is_heading": is_heading,
                "alignment": str(para.alignment) if para.alignment else None,
                "paragraph_index": len(paragraphs)
            })
        return paragraphs

    def extract_styles(self) -> Dict[str, Any]:
        """
        返回文档中的样式信息

        Returns:
            Dict: 样式名称到样式信息的映射
        """
        styles_info = {}
        for style in self.document.styles:
            if style.type == 1:  # Paragraph style
                styles_info[style.name] = {
                    "name": style.name,
                    "type": "paragraph",
                    "font_name": None,
                    "font_size": None,
                }
                try:
                    if style.font:
                        styles_info[style.name]["font_name"] = style.font.name
                        if style.font.size:
                            styles_info[style.name]["font_size"] = style.font.size.pt
                except Exception:
                    pass
        return styles_info

    def extract_headers_footers(self) -> Dict[str, List[str]]:
        """
        提取页眉页脚内容

        Returns:
            Dict: 包含 headers 和 footers 的字典
        """
        result = {
            "headers": [],
            "footers": []
        }

        # 遍历所有节获取页眉页脚
        for section in self.document.sections:
            # 获取页眉
            for header in section.header.paragraphs:
                if header.text.strip():
                    result["headers"].append(header.text.strip())

            # 获取页脚
            for footer in section.footer.paragraphs:
                if footer.text.strip():
                    result["footers"].append(footer.text.strip())

        return result

    def extract_headings(self) -> List[Tuple[int, str, int]]:
        """
        返回标题结构

        Returns:
            List[Tuple]: [(level, text, paragraph_index), ...]
        """
        if self._headings_cache is not None:
            return self._headings_cache

        headings = []
        for idx, para in enumerate(self.document.paragraphs):
            text = para.text.strip()
            style_name = para.style.name if para.style else "Normal"
            level = self._get_heading_level(style_name)

            if level is not None and text:
                headings.append((level, text, idx))

        self._headings_cache = headings
        return headings

    def detect_language(self) -> str:
        """
        检测文档语言

        Returns:
            str: "chinese" 如果中文比例 > 30%，否则 "english"
        """
        total_chars = 0
        chinese_chars = 0

        # 中文字符范围
        chinese_pattern = re.compile(r'[\u4e00-\u9fff]')

        for para in self.document.paragraphs:
            text = para.text
            total_chars += len(text)
            chinese_chars += len(chinese_pattern.findall(text))

        if total_chars == 0:
            return "english"

        chinese_ratio = chinese_chars / total_chars
        return "chinese" if chinese_ratio > 0.3 else "english"

    def get_page_setup(self) -> Dict[str, Any]:
        """
        返回页面设置

        Returns:
            Dict: 包含纸张大小、页边距等设置
        """
        if not self.document.sections:
            return {}

        section = self.document.sections[0]
        page_width = section.page_width
        page_height = section.page_height

        # 转换为厘米
        page_width_cm = page_width.cm if page_width else 21.0
        page_height_cm = page_height.cm if page_height else 29.7

        return {
            "paper_width_cm": round(page_width_cm, 2),
            "paper_height_cm": round(page_height_cm, 2),
            "margin_top_cm": round(section.top_margin.cm, 2),
            "margin_bottom_cm": round(section.bottom_margin.cm, 2),
            "margin_left_cm": round(section.left_margin.cm, 2),
            "margin_right_cm": round(section.right_margin.cm, 2),
            "orientation": "landscape" if page_width_cm > page_height_cm else "portrait"
        }

    def _is_heading_style(self, style_name: str) -> bool:
        """判断是否为标题样式"""
        heading_keywords = ["Heading", "标题", "TOC"]
        return any(keyword.lower() in style_name.lower() for keyword in heading_keywords)

    def _get_heading_level(self, style_name: str) -> int | None:
        """
        获取标题级别

        Returns:
            int: 标题级别 (1-9)，None 表示不是标题
        """
        style_lower = style_name.lower()

        # 处理 "Heading 1", "Heading1" 等格式
        if "heading" in style_lower or "标题" in style_lower:
            # 尝试提取数字
            match = re.search(r'(\d+)', style_name)
            if match:
                return int(match.group(1))

            # 中文标题格式
            match = re.search(r'标题[一二三四五六七八九]', style_name)
            if match:
                chinese_levels = {'一': 1, '二': 2, '三': 3, '四': 4, '五': 5,
                                  '六': 6, '七': 7, '八': 8, '九': 9}
                return chinese_levels.get(match.group(0)[-1])

        return None

    def get_paragraph_format(self, paragraph_index: int) -> Dict[str, Any]:
        """
        获取指定段落的格式信息

        Args:
            paragraph_index: 段落索引

        Returns:
            Dict: 格式信息
        """
        if paragraph_index >= len(self.document.paragraphs):
            return {}

        para = self.document.paragraphs[paragraph_index]
        pf = para.paragraph_format

        return {
            "first_line_indent": pf.first_line_indent.pt if pf.first_line_indent else 0,
            "left_indent": pf.left_indent.pt if pf.left_indent else 0,
            "right_indent": pf.right_indent.pt if pf.right_indent else 0,
            "space_before": pf.space_before.pt if pf.space_before else 0,
            "space_after": pf.space_after.pt if pf.space_after else 0,
            "line_spacing": pf.line_spacing if pf.line_spacing else None,
        }

    def get_run_font_info(self, paragraph_index: int) -> List[Dict[str, Any]]:
        """
        获取段落中所有 run 的字体信息

        Args:
            paragraph_index: 段落索引

        Returns:
            List[Dict]: 每个 run 的字体信息
        """
        if paragraph_index >= len(self.document.paragraphs):
            return []

        para = self.document.paragraphs[paragraph_index]
        runs_info = []

        for run in para.runs:
            font = run.font
            runs_info.append({
                "text": run.text,
                "font_name": font.name,
                "font_size": font.size.pt if font.size else None,
                "bold": font.bold,
                "italic": font.italic,
                "underline": font.underline,
            })

        return runs_info
